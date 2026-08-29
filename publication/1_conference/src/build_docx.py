"""Rebuild the ISGT Asia draft (Word) from the pristine IEEE template.

Design rules:
  * Start from the PRISTINE backup so the IEEE 2-column section layout survives.
  * Full prose, not notes: four numbered sections, abstract and index terms.
  * HARD LIMIT of four pages including references. `preview.py` renders the
    built file through Word and reports the true page count - python-docx
    cannot paginate, so that is the only way to know. Several choices here
    exist only to hold that limit: five figures rather than seven, one table
    rather than four, body space-after removed, equation leading halved.
  * Citations are IEEE numeric, numbered by order of first use in the text
    (see cite_number). The field result is plain text so the document reads
    without Zotero, and the CSL payload lets Zotero re-render it; the document
    preferences declare the IEEE style. A matching .ris is exported alongside.
  * Tables use the journal three-line rule and are kept whole across columns.
  * Equations are native Word equation objects (OMML), right-numbered (1)..(6),
    with notation defined in the sentence around them rather than in a
    detached where-clause.
  * Wide figures sit in full-width (single-column) islands created with
    continuous section breaks, IEEE "figure*" style. Note that text placed
    BEFORE such an island cannot shorten the document tail - freeing it only
    widens the gap at the foot of the preceding page.

Run:  python build_docx.py            # writes to $ISGT_DOC_DIR or OneDrive
      python preview.py <docx> <dir>  # page count + page images
"""

from __future__ import annotations

import copy
import datetime
import json
import os
import re
import shutil
from pathlib import Path

import docx
import pandas as pd
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

import config
import references as refs

# --- Paths ----------------------------------------------------------------
DOC_DIR = Path(os.environ.get("ISGT_DOC_DIR") or (
    r"C:\Users\z5404477\OneDrive - UNSW\H0424909\06_Seminar Conference etc"
    r"\2026-10 ISGT Asia\2 - Silver"
))
DOCX_PATH = DOC_DIR / "ISGT Asia - SS - Load Backbone Estimation.docx"
# The pristine IEEE template is vendored into the repo. It used to be read from
# a timestamped backup next to the draft, which broke as soon as those backups
# were tidied into an archive folder.
TEMPLATE_PATH = Path(os.environ.get("ISGT_TEMPLATE")
                     or config.ROOT / "assets" / "ieee_template.docx")
RIS_PATH = DOC_DIR / "AQF_paper_references.ris"

FIG = config.GOLD_FIGURES_DIR
TABLES_DIR = config.GOLD_TABLES_DIR
OMML_PATH = Path(__file__).with_name("equations_omml.json")

TITLE = ("Adaptive Quantile Flexibility: Recovering Candidate Flexible Load "
         "from Aggregate Profiles for Network Tariff Modelling")

# Derived from config so the width a figure is AUTHORED at (plotting.py) and the
# width it is INSERTED at can never drift apart. If they drift, Word rescales the
# image and every point size inside the figure changes silently.
COL_W = Inches(config.FIG_COL_W_IN)    # one IEEE column
FULL_W = Inches(config.FIG_FULL_W_IN)  # both columns
TABLE_PT = 8           # matches the template's table/caption size
M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"
ZOTERO_LIB = "http://zotero.org/users/local/aqfIsgt26/items/"


# =========================================================================
# Low-level OXML helpers
# =========================================================================
def backup(path: Path) -> Path:
    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    dest = path.with_name(f"{path.stem}_backup_{ts}{path.suffix}")
    shutil.copy(path, dest)
    return dest


def fld_char(kind: str):
    el = OxmlElement("w:fldChar")
    el.set(qn("w:fldCharType"), kind)
    return el


def instr_text(text: str):
    el = OxmlElement("w:instrText")
    el.set(XML_SPACE, "preserve")
    el.text = text
    return el


def omml_element(xml: str):
    """Parse a bare <m:oMath> string into an element with namespaces bound."""
    wrapped = f'<root xmlns:m="{M_NS}" xmlns:w="{W_NS}">{xml}</root>'
    return parse_xml(wrapped)[0]


def set_mark_size(paragraph, half_points: int):
    """Shrink an empty paragraph's own mark (its height) without touching text."""
    pPr = paragraph._p.get_or_add_pPr()
    rPr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(half_points))
    rPr.append(sz)
    sect = pPr.find(qn("w:sectPr"))
    if sect is not None:                       # rPr must precede sectPr
        sect.addprevious(rPr)
    else:
        pPr.append(rPr)


def unnumber(paragraph):
    """Strip a style's automatic list numbering from one paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numId = OxmlElement("w:numId")
    numId.set(qn("w:val"), "0")
    numPr.append(ilvl)
    numPr.append(numId)
    pPr.insert(0 if pPr.find(qn("w:pStyle")) is None else 1, numPr)


# OOXML requires a strict child order inside w:tblPr; appending silently
# produces a file Word refuses to open ("unreadable content").
TBLPR_ORDER = [
    "tblStyle", "tblpPr", "tblOverlap", "bidiVisual", "tblStyleRowBandSize",
    "tblStyleColBandSize", "tblW", "jc", "tblCellSpacing", "tblInd", "tblBorders",
    "shd", "tblLayout", "tblCellMar", "tblLook", "tblCaption", "tblDescription",
]


def insert_ordered(parent, element, order):
    """Place `element` at its schema-mandated position among `parent`'s children."""
    name = element.tag.split("}")[1]
    rank = order.index(name)
    for child in parent:
        child_name = child.tag.split("}")[1]
        if child_name in order and order.index(child_name) > rank:
            child.addprevious(element)
            return element
    parent.append(element)
    return element


def set_three_line_borders(table, top=8, mid=4, bottom=8):
    """Journal 'booktabs' rules: top rule, rule under header row, bottom rule.
    Sizes are eighths of a point (8 = 1.0 pt, 4 = 0.5 pt). No other lines."""
    last = len(table.rows) - 1
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            for old in tcPr.findall(qn("w:tcBorders")):
                tcPr.remove(old)
            borders = OxmlElement("w:tcBorders")
            spec = [
                ("top", top if r_idx == 0 else 0),
                ("left", 0),
                ("bottom", mid if r_idx == 0 else (bottom if r_idx == last else 0)),
                ("right", 0),
                ("insideH", 0),
                ("insideV", 0),
            ]
            for name, sz in spec:                       # order matters in OOXML
                el = OxmlElement(f"w:{name}")
                if sz:
                    el.set(qn("w:val"), "single")
                    el.set(qn("w:sz"), str(sz))
                    el.set(qn("w:space"), "0")
                    el.set(qn("w:color"), "000000")
                else:
                    el.set(qn("w:val"), "nil")
                borders.append(el)
            tcPr.append(borders)


def set_fixed_layout(table, widths):
    """Pin the table to an exact width and column grid, so Word cannot re-flow
    (or silently drop) columns inside a narrow IEEE column."""
    tblPr = table._tbl.tblPr
    for tag in ("w:tblW", "w:tblLayout"):
        for el in tblPr.findall(qn(tag)):
            tblPr.remove(el)
    w_el = OxmlElement("w:tblW")
    w_el.set(qn("w:w"), str(sum(int(w.inches * 1440) for w in widths)))
    w_el.set(qn("w:type"), "dxa")
    insert_ordered(tblPr, w_el, TBLPR_ORDER)
    lay = OxmlElement("w:tblLayout")
    lay.set(qn("w:type"), "fixed")
    insert_ordered(tblPr, lay, TBLPR_ORDER)
    grid = table._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for gc, width in zip(grid.findall(qn("w:gridCol")), widths):
            gc.set(qn("w:w"), str(int(width.inches * 1440)))
    trPr = table.rows[0]._tr.get_or_add_trPr()          # repeat header on break
    trPr.append(OxmlElement("w:tblHeader"))


def section_break(anchor, base_sectpr, num_cols: int, space: int):
    """Insert an almost-invisible paragraph carrying a continuous section break."""
    p = anchor.insert_paragraph_before()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    sect = copy.deepcopy(base_sectpr)
    typ = sect.find(qn("w:type"))
    if typ is None:
        typ = OxmlElement("w:type")
        sect.insert(0, typ)
    typ.set(qn("w:val"), "continuous")
    cols = sect.find(qn("w:cols"))
    if num_cols == 1:
        cols.attrib.pop(qn("w:num"), None)
    else:
        cols.set(qn("w:num"), str(num_cols))
    cols.set(qn("w:space"), str(space))
    p._p.get_or_add_pPr().append(sect)                  # sectPr must be last
    set_mark_size(p, 4)
    return p


# =========================================================================
# Zotero-linked citations
# =========================================================================
_cite_counter = [0]


def _item_payload(key: str, idx: int) -> dict:
    rec = refs.REFS[key]
    uri = ZOTERO_LIB + f"{abs(hash(key)) % 10**8:08d}"
    item = dict(rec["csl"])
    item["id"] = 1000 + idx
    return {"id": uri, "uris": [uri], "itemData": item}


# IEEE numbers references by order of FIRST CITATION in the text, not
# alphabetically. Numbers are therefore assigned as the body is built, and the
# reference list is emitted afterwards in this order.
_cite_order: list[str] = []


def cite_number(key: str) -> int:
    """1-based IEEE reference number, assigned on first use."""
    if key not in _cite_order:
        _cite_order.append(key)
    return _cite_order.index(key) + 1


def _format_numeric(keys: list[str]) -> str:
    """One bracket per citation group: '[3]', '[3, 7]', '[3-5]', '[3-5, 9]'.

    Grouped rather than the IEEE '[3], [7]' form, at the author's request. The
    document's Zotero style is set to match (see main), so refreshing the fields
    reproduces this instead of reverting.
    """
    nums = sorted(cite_number(k) for k in keys)
    groups: list[list[int]] = []
    for n in nums:
        if groups and n == groups[-1][-1] + 1:
            groups[-1].append(n)
        else:
            groups.append([n])
    parts = [str(g[0]) if len(g) == 1 else
             (f"{g[0]}, {g[1]}" if len(g) == 2 else f"{g[0]}\u2013{g[-1]}")
             for g in groups]
    return "[" + ", ".join(parts) + "]"


def add_citation(paragraph, keys: list[str], prefix: str = " "):
    """One Zotero citation field, displayed as an IEEE bracketed number.

    The field result is plain text, so the document reads correctly for anyone
    without Zotero; the CSL payload lets Zotero re-render the whole bibliography
    (in IEEE style, per the document preferences written at the end of main).
    """
    _cite_counter[0] += 1
    display = _format_numeric(keys)
    blob = {
        "citationID": f"aqf{_cite_counter[0]:04d}",
        "properties": {"formattedCitation": display, "plainCitation": display, "noteIndex": 0},
        "citationItems": [_item_payload(k, i) for i, k in enumerate(sorted(keys))],
        "schema": "https://github.com/citation-style-language/schema/raw/master/csl-citation.json",
    }
    instr = " ADDIN ZOTERO_ITEM CSL_CITATION " + json.dumps(blob, ensure_ascii=False) + " "
    if prefix:
        paragraph.add_run(prefix)
    paragraph.add_run()._r.append(fld_char("begin"))
    paragraph.add_run()._r.append(instr_text(instr))
    paragraph.add_run()._r.append(fld_char("separate"))
    paragraph.add_run(display)
    paragraph.add_run()._r.append(fld_char("end"))


CITE_RE = re.compile(r"\[\[([^\]]+)\]\]")

# =========================================================================
# Word fields: SEQ numbering for figures/tables/equations, REF for xrefs
# =========================================================================
# Numbers are Word fields rather than typed text, so inserting a figure or an
# equation in the middle renumbers everything on refresh (Ctrl+A, F9). Each
# field is written with a cached result as well, so the document reads correctly
# before anyone refreshes it.
_bookmark_id = [2000]

# Document order, fixed up front so prose can reference a figure before the
# figure is inserted.
FIGURE_ORDER = ["decomp", "positioning", "mechanism", "map", "curves"]
TABLE_ORDER = ["results"]
EQUATION_ORDER = ["mvp", "model", "relation", "qstar", "guard", "metrics"]

_ROMAN = ["I", "II", "III", "IV", "V", "VI", "VII", "VIII"]


def _bookmark_name(kind: str, key: str) -> str:
    return f"_{kind}_{key}"


def figure_number(key: str) -> int:
    return FIGURE_ORDER.index(key) + 1


def table_number(key: str) -> str:
    return _ROMAN[TABLE_ORDER.index(key)]


def equation_number(key: str) -> int:
    return EQUATION_ORDER.index(key) + 1


def add_field(paragraph, instr: str, cached: str, *, size_pt=None, bold=None):
    """Append a Word field, carrying a cached result. Returns (first, last) runs."""
    made = []

    def _run(child):
        r = OxmlElement("w:r")
        if size_pt is not None or bold is not None:
            rPr = OxmlElement("w:rPr")
            if bold:
                rPr.append(OxmlElement("w:b"))
            if size_pt is not None:
                sz = OxmlElement("w:sz")
                sz.set(qn("w:val"), str(int(size_pt * 2)))
                rPr.append(sz)
            r.append(rPr)
        r.append(child)
        paragraph._p.append(r)
        made.append(r)
        return r

    _run(fld_char("begin"))
    _run(instr_text(instr))
    _run(fld_char("separate"))
    t = OxmlElement("w:t")
    t.text = cached
    _run(t)
    _run(fld_char("end"))
    return made[0], made[-1]


def wrap_bookmark(name: str, first_el, last_el):
    _bookmark_id[0] += 1
    bid = str(_bookmark_id[0])
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), bid)
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bid)
    first_el.addprevious(start)
    last_el.addnext(end)


def seq_field(paragraph, kind: str, key: str, seq_name: str, fmt: str, cached: str,
              *, size_pt=None, bold=None):
    """A bookmarked { SEQ } field, so { REF } elsewhere can point at it."""
    first, last = add_field(paragraph, f" SEQ {seq_name} \\* {fmt} ", cached,
                            size_pt=size_pt, bold=bold)
    wrap_bookmark(_bookmark_name(kind, key), first, last)
    return paragraph


def ref_field(paragraph, kind: str, key: str, cached: str):
    """A { REF } cross-reference to a bookmarked SEQ field."""
    return add_field(paragraph, f" REF {_bookmark_name(kind, key)} \\h ", cached)



# Math symbols written as ASCII in the source (R_F, MAE_B, q_def, D_th) are set
# as real subscripts rather than printed with a literal underscore. The base is
# italicised only when it is a single letter: multi-letter operators such as MAE
# are upright by convention.
# Base may carry a combining circumflex (D-hat, kappa-hat); w, p and kappa are
# also legitimate subscripted symbols in the prose.
SUBSCRIPT_RE = re.compile("\\b(MAE|[REBFqDwp\u03ba]\u0302?)_([A-Za-z0-9]{1,3})\\b")

# Subscripts that are running indices are italic maths; descriptive subscripts
# (F for flexible, def for default, th for threshold, day) stay upright.
INDEX_SUBS = {"t", "i", "j", "n"}

# IEEE sets every scalar variable in italic, in running text as well as in the
# displayed equations. Prose marks those spans as $x$; the delimiters never
# reach the page. Only letters go inside them - digits and operators stay
# upright, so write "(1-$p$)/2", not "$(1-p)/2$".
MATH_RE = re.compile(r"\$([^$]+)\$")


XREF_RE = re.compile(r"\{\{(fig|tab|eq):([a-z_]+)\}\}")


def add_plain_text(paragraph, text: str):
    """Add text, expanding {{fig:key}}-style cross-references, $x$ math spans
    and NAME_SUB tokens (rendered with a true subscript run)."""
    pos = 0
    for m in XREF_RE.finditer(text):
        if m.start() > pos:
            _add_math(paragraph, text[pos:m.start()])
        kind, key = m.group(1), m.group(2)
        cached = {"fig": lambda: str(figure_number(key)),
                  "tab": lambda: table_number(key),
                  "eq": lambda: str(equation_number(key))}[kind]()
        ref_field(paragraph, kind, key, cached)
        pos = m.end()
    if pos < len(text):
        _add_math(paragraph, text[pos:])


def _add_math(paragraph, text: str):
    """Split on $...$ spans; everything inside one is italic maths."""
    pos = 0
    for m in MATH_RE.finditer(text):
        if m.start() > pos:
            _add_subscripted(paragraph, text[pos:m.start()], italic=False)
        _add_subscripted(paragraph, m.group(1), italic=True)
        pos = m.end()
    if pos < len(text):
        _add_subscripted(paragraph, text[pos:], italic=False)


def _add_subscripted(paragraph, text: str, italic: bool = False):
    pos = 0
    for m in SUBSCRIPT_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()]).italic = italic
        base, sub = m.group(1), m.group(2)
        base_run = paragraph.add_run(base)
        # Single-letter bases (hatted or not) are italic maths; multi-letter
        # operators such as MAE stay upright.
        base_run.italic = len(base.replace("̂", "")) == 1
        sub_run = paragraph.add_run(sub)
        sub_run.font.subscript = True
        sub_run.italic = sub in INDEX_SUBS
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:]).italic = italic


def add_rich_text(paragraph, text: str):
    """Text with [[key|key]] citation markers turned into Zotero fields."""
    pos = 0
    for m in CITE_RE.finditer(text):
        if m.start() > pos:
            add_plain_text(paragraph, text[pos:m.start()])
        add_citation(paragraph, m.group(1).split("|"), prefix="")
        pos = m.end()
    if pos < len(text):
        add_plain_text(paragraph, text[pos:])
    return paragraph


# =========================================================================
# Content helpers
# =========================================================================
def bullets(anchor, texts):
    for t in texts:
        add_rich_text(anchor.insert_paragraph_before(style="bullet list"), t)


def picture(anchor, key: str, image_name: str, caption: str, width):
    """Figure plus a caption numbered by a { SEQ Figure } field.

    The style's own list numbering is stripped: a SEQ field can be bookmarked
    and cross-referenced, list numbering cannot.
    """
    p = anchor.insert_paragraph_before()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.keep_with_next = True    # image stays with caption
    p.add_run().add_picture(str(FIG / image_name), width=width)

    cap = anchor.insert_paragraph_before(style="figure caption")
    unnumber(cap)
    cap.add_run("Fig. ")
    seq_field(cap, "fig", key, "Figure", "ARABIC", str(figure_number(key)))
    cap.add_run(".\u2003")
    add_plain_text(cap, caption)
    return cap


def table_caption(anchor, key: str, caption: str):
    """Table caption numbered by a { SEQ Table \\* ROMAN } field."""
    p = anchor.insert_paragraph_before(style="table head")
    unnumber(p)
    p.add_run("Table ")
    seq_field(p, "tab", key, "Table", "ROMAN", table_number(key))
    p.add_run(".\u2003")
    add_plain_text(p, caption)
    p.paragraph_format.keep_with_next = True    # caption must not orphan
    return p


_ALIGN = {
    "l": WD_ALIGN_PARAGRAPH.LEFT,
    "c": WD_ALIGN_PARAGRAPH.CENTER,
    "r": WD_ALIGN_PARAGRAPH.RIGHT,
}


def keep_table_together(table):
    """Stop a short table from breaking across a column or page.

    Word will otherwise split a four-row table at a column boundary and repeat
    the header on the far side, which reads as two tables. Every row but the
    last is marked keep-with-next, and no row may split internally.
    """
    for row in table.rows[:-1]:
        # w:cantSplit must precede the w:tblHeader that set_fixed_layout adds.
        insert_ordered(row._tr.get_or_add_trPr(), OxmlElement("w:cantSplit"),
                       _PR_ORDERS["trPr"])
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.keep_with_next = True
    for cell in table.rows[-1].cells:
        for para in cell.paragraphs:
            para.paragraph_format.keep_with_next = False


def dataframe_table(doc, anchor, df: pd.DataFrame, widths, font_pt=TABLE_PT,
                    align=None, headers=None, bold_cells=(), italic_rows=()):
    """Render a DataFrame as an IEEE three-line table.

    align       per-column "l"/"c"/"r"; defaults to a left label column and
                centred remainder. Numeric columns should be pre-formatted to a
                fixed number of decimals by the caller so they line up.
    headers     verbatim header strings. Supply these whenever a header contains
                mathematics - the automatic title-casing below mangles anything
                that is not plain words.
    bold_cells  {(row, col)} zero-based DATA indices to bold, for marking the
                best value in a column.
    italic_rows {row} zero-based DATA indices to italicise, for rows that are a
                reference point rather than a competitor.
    """
    table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
    table.autofit = False

    if align is None:
        align = ["l"] + ["c"] * (df.shape[1] - 1)

    def head_case(text):
        return " ".join(w.capitalize() if w.isalpha() and w.islower() else w
                        for w in str(text).split())

    if headers is None:
        # underscores joining two lowercase words are separators; keep R_F, MAE_B
        headers = [head_case(re.sub(r"(?<=[a-z])_(?=[a-z])", " ", str(c)))
                   for c in df.columns]
    for j, head in enumerate(headers):
        cell = table.rows[0].cells[j]
        para = cell.paragraphs[0]
        # Built run-by-run rather than via cell.text, so NAME_SUB tokens in a
        # header become real subscripts like the body text does.
        for k, line in enumerate(str(head).splitlines() or [""]):
            if k:
                para.add_run()._r.append(OxmlElement("w:br"))
            add_plain_text(para, line)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.bold = True
            run.font.size = Pt(font_pt)
    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(df.iat[i, j])
            cell.paragraphs[0].alignment = _ALIGN[align[j]]
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(font_pt)
                if (i, j) in bold_cells:
                    run.bold = True
                if i in italic_rows:
                    run.italic = True
    for j, w in enumerate(widths):                      # width must be set per cell
        for row in table.rows:
            row.cells[j].width = w
    set_fixed_layout(table, widths)
    set_three_line_borders(table)
    keep_table_together(table)
    anchor._p.addprevious(table._tbl)
    return table


def equation(anchor, omml_xml: str, key: str):
    """A display equation, right-numbered by a { SEQ Equation } field.

    Every symbol is defined in the sentence that introduces the equation rather
    than in a detached where-clause: it reads better and costs far less space.
    """
    p = anchor.insert_paragraph_before(style="equation")
    r1 = p.add_run()
    r1._r.append(OxmlElement("w:tab"))
    p._p.append(omml_element(omml_xml))
    r2 = p.add_run()
    r2._r.append(OxmlElement("w:tab"))
    r2.add_text("(")
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(10)
    seq_field(p, "eq", key, "Equation", "ARABIC", str(equation_number(key)), size_pt=10)
    r3 = p.add_run(")")
    r3.font.name = "Times New Roman"
    r3.font.size = Pt(10)
    return p


def body(anchor, texts):
    """Justified body paragraphs, with [[key]] markers expanded to Zotero fields."""
    for t in texts:
        add_rich_text(anchor.insert_paragraph_before(style="Body Text"), t)


def heading(anchor, text: str):
    """A numbered section heading (the style supplies the roman numeral)."""
    return anchor.insert_paragraph_before(text, style="Heading 1")


def subheading(anchor, text: str):
    """A lettered subsection heading under a Methodology-style section."""
    return anchor.insert_paragraph_before(text, style="Heading 2")


def find_tail_anchor(doc):
    """The paragraph carrying the 2-column sectPr. Deleting it collapses the
    whole paper to a single column - it is the end anchor, never content."""
    for p in doc.paragraphs:
        sect = p._p.find(qn("w:pPr") + "/" + qn("w:sectPr"))
        if sect is not None:
            cols = sect.find(qn("w:cols"))
            if cols is not None and cols.get(qn("w:num")) == "2":
                return p, sect
    raise RuntimeError("2-column section break paragraph not found in template")


def clear_between(doc, start, end):
    body = doc.element.body
    children = list(body)
    for child in children[children.index(start._p) + 1: children.index(end._p)]:
        body.remove(child)


def fix_equation_font(doc):
    """Stop display equations rendering their function names as Greek.

    The template's "equation" style is set in the Symbol font. OMML runs marked
    <m:nor/> - which is how multi-letter function names are kept upright - take
    the paragraph font rather than Cambria Math, and Symbol maps ASCII onto
    Greek: e->epsilon, r->rho, n->nu, m->mu, a->alpha, x->xi, d->delta,
    f->phi, t->tau, h->eta. So "Bern" printed as "Bern" with Greek letters,
    "min" as a mu-iota-nu string, "max" likewise, and "q_def"/"D_th" the same.
    Setting the style to Cambria Math renders them as written; ordinary math
    runs are unaffected because Word already sets those in Cambria Math.
    """
    style = doc.styles["equation"]
    style.font.name = "Cambria Math"
    rpr = style.element.find(qn("w:rPr"))
    if rpr is not None:
        fonts = rpr.find(qn("w:rFonts"))
        if fonts is not None:
            for attr in ("w:ascii", "w:hAnsi", "w:cs"):
                fonts.set(qn(attr), "Cambria Math")


def tighten_styles(doc):
    """Trim the template's inter-paragraph leading to IEEE-conventional values.

    The template ships 6 pt after every body paragraph and 12 pt on both sides of
    every display equation. Across 28 body paragraphs and 8 equations that is
    close to seven column-inches - most of a page against a hard four-page limit.
    IEEE body copy separates paragraphs by first-line indent alone, so removing
    the body space-after is the standard look as well as the cheaper one; the
    equations keep half their original leading.
    """
    body = doc.styles["Body Text"].paragraph_format
    body.space_after = Pt(0)
    eq = doc.styles["equation"].paragraph_format
    eq.space_before = Pt(6)
    eq.space_after = Pt(6)
    cap = doc.styles["figure caption"].paragraph_format
    cap.space_after = Pt(6)
    head = doc.styles["table head"].paragraph_format
    head.space_before = Pt(8)
    doc.styles["references"].paragraph_format.space_after = Pt(0)


def kill_autospacing(paragraph):
    """Clear w:beforeAutospacing / w:afterAutospacing on a paragraph.

    These are the reason the author block floats: Word expands each flagged
    empty paragraph to roughly 14 pt of leading no matter what point size the
    paragraph mark carries, so shrinking the mark alone does nothing.
    """
    pPr = paragraph._p.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        # w:spacing must precede w:rPr and w:sectPr, so it cannot just be appended.
        spacing = insert_ordered(pPr, OxmlElement("w:spacing"), _PR_ORDERS["pPr"])
    for attr in ("w:beforeAutospacing", "w:afterAutospacing"):
        spacing.set(qn(attr), "0")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")


def tighten_author_block(doc, first_heading):
    """Collapse the empty spacer paragraphs around the two author rows.

    A paragraph between two adjacent tables is structurally required - remove it
    and Word merges the tables into one - so spacers that sit between tables, or
    that carry the section properties, are collapsed rather than deleted.
    """
    body_children = list(doc.element.body)
    stop = body_children.index(first_heading._p)
    removable = []
    for p in doc.paragraphs:
        idx = body_children.index(p._p)
        if idx >= stop:
            break
        if p.text.strip():
            continue
        has_sect = p._p.find(qn("w:pPr") + "/" + qn("w:sectPr")) is not None
        neighbours = {body_children[idx - 1].tag.split("}")[1] if idx else "",
                      body_children[idx + 1].tag.split("}")[1]}
        if has_sect or "tbl" in neighbours:
            set_mark_size(p, 4)                 # 2 pt mark
            kill_autospacing(p)                 # ... and no auto leading around it
        else:
            removable.append(p)
    for p in removable:                          # pure spacers: drop entirely
        p._p.getparent().remove(p._p)


def build_results_summary_table() -> pd.DataFrame:
    """Legacy mean-only summary.

    Kept so table3_results_summary.csv stays in step with the data, but the
    paper no longer prints it: ranking by the mean of |R_F - 1| is dominated by
    the low-persistence cells where the truncation noise floor, not the
    estimator, inflates R_F. Table III (build_headline_table) replaces it.
    """
    gold = pd.read_csv(TABLES_DIR / "recoverability_summary.csv")
    rows = []
    for variant, label in [
        ("fixed_q_0.1", "Fixed q=0.1"), ("fixed_q_0.2", "Fixed q=0.2"),
        ("fixed_q_0.3", "Fixed q=0.3"), ("oracle_aqf", "Oracle-AQF"),
        ("estimated_aqf", "Estimated-AQF"),
    ]:
        sub = gold[gold["estimator_variant"] == variant]
        rows.append({"estimator": label,
                     "Mean R_F": round(sub["r_f_mean"].mean(), 2),
                     "Mean |R_F \u2212 1|": round((sub["r_f_mean"] - 1).abs().mean(), 2)})
    df = pd.DataFrame(rows).sort_values("Mean |R_F \u2212 1|").reset_index(drop=True)
    df.to_csv(TABLES_DIR / "table3_results_summary.csv", index=False)
    return df


ORACLE_LABEL = "Oracle-AQF"          # key used in the Gold CSV
# Printed names. "Oracle" is jargon for a non-specialist reader, so the row is
# labelled by what it actually is: a bound computed with the true values.
DISPLAY_LABELS = {
    "Oracle-AQF": "Oracle-q",
    "Estimated-AQF": "AQF",
    "Fixed q=0.3": "Fixed q = 0.3",
    "Fixed q=0.2": "Fixed q = 0.2",
    "Fixed q=0.1": "Fixed q = 0.1",
}


def build_headline_table() -> tuple[pd.DataFrame, list, set, set]:
    """Table III: the four numbers the Results section actually argues from.

    Returns the printable frame plus the headers, the cells to bold, and the
    rows to italicise. Oracle-q knows the true p and kappa, so it is a reference
    rather than a competitor: its row is italicised and excluded from the
    best-in-column bolding. It is NOT a performance ceiling - it still takes an
    empirical quantile of D days, so AQF beats it on backbone error wherever the
    signal is strong.
    """
    src = pd.read_csv(TABLES_DIR / "table3_headline_results.csv")
    # Stacked headers: five columns have to fit one 3.3-inch IEEE column.
    cols = [
        ("mean_abs_rf_dev", "Mean\n|R_F \u2212 1|"),
        ("median_abs_rf_dev", "Median\n|R_F \u2212 1|"),
        ("median_abs_rf_dev_identifiable",
         f"Median,\n$\u03ba$ \u2265 {config.KAPPA_IDENTIFIABLE:.2f}"),
        ("mean_mae_b", "MAE_B\n(kW)"),
    ]
    out = pd.DataFrame({"estimator": src["estimator"].map(
        lambda v: DISPLAY_LABELS.get(v, v))})
    for key, _ in cols:
        out[key] = src[key].map(lambda v: f"{v:.3f}")   # fixed decimals, so columns align

    deployable = src["estimator"] != ORACLE_LABEL
    bold = set()
    for j, (key, _) in enumerate(cols, start=1):
        best = src.loc[deployable, key].idxmin()
        bold.add((int(best), j))
    italic = {int(i) for i in src.index[~deployable]}
    headers = ["Estimator"] + [h for _, h in cols]
    return out, headers, bold, italic


def build_win_count_table() -> pd.DataFrame:
    """Table IV: head-to-head cell wins for estimated-AQF.

    All four opponents are reported together on purpose. A practitioner cannot
    know in advance which fixed q suits a given timestamp, so the per-cell best
    fixed q is itself an oracle; quoting only the most favourable comparison
    would be baseline shopping.
    """
    src = pd.read_csv(TABLES_DIR / "table4_win_counts.csv")
    return pd.DataFrame({
        "opponent": src["opponent"],
        "available": src["available_in_advance"],
        "won": [f"{w} / {n}" for w, n in zip(src["cells_won"], src["cells_total"])],
    })


# =========================================================================
# Front matter
# =========================================================================
ABSTRACT = (
    "Network businesses are beginning to use tariffs to move demand in time: the "
    "Solar Sharer Offer in New South Wales, for example, gives households free "
    "electricity in the middle of the day. Sizing the effect of such an offer "
    "first requires knowing how much of the load at each time of day could "
    "actually move. In the models utilities can run today it is assumed rather "
    "than measured – "
    "either as a single shiftable fraction of daily energy applied everywhere, or "
    "through a bottom-up appliance model needing data that utilities do not hold. The "
    "practical middle ground takes the non-shiftable \u201cbackbone\u201d to be a fixed low "
    "quantile of the loads recorded at the same time of day across many days, but "
    "that quantile is hand-picked and one value is applied to every time of day. "
    "This paper shows the correct quantile is not a constant: it depends on how "
    "often flexible events occur and how large they are next to ordinary "
    "day-to-day variation. We derive the relation between those quantities in "
    "closed form and use it to build Adaptive Quantile Flexibility (AQF), which "
    "computes the quantile separately at each timestamp from the meter data "
    "alone, declares when the data cannot support that estimate and falls back to "
    "a default quantile, and never exceeds the median. Tested on synthetic load "
    "profiles where the true split is known by construction, AQF more than halves "
    "the median error in recovered flexible energy relative to the best fixed "
    "quantile and cuts backbone error roughly threefold."
)

KEYWORDS = ("demand flexibility, demand response, load disaggregation, quantile "
            "estimation, smart metering, tariff design")


# =========================================================================
# Post-save structural check.  Word refuses to open a .docx whose XML breaks
# the OOXML content model, and it gives no useful diagnostic, so verify here.
# =========================================================================
_P_OK = {"pPr", "r", "hyperlink", "fldSimple", "subDoc", "customXml", "smartTag", "sdt",
         "dir", "bdo", "bookmarkStart", "bookmarkEnd", "commentRangeStart",
         "commentRangeEnd", "ins", "del", "moveFrom", "moveTo", "proofErr",
         "permStart", "permEnd", "oMath", "oMathPara"}
_R_OK = {"rPr", "t", "br", "tab", "fldChar", "instrText", "delText", "noBreakHyphen",
         "softHyphen", "sym", "cr", "drawing", "object", "pict", "footnoteReference",
         "endnoteReference", "commentReference", "ptab", "lastRenderedPageBreak",
         "ruby", "separator", "continuationSeparator", "annotationRef", "footnoteRef",
         "endnoteRef", "AlternateContent"}
_PR_ORDERS = {
    "pPr": ["pStyle", "keepNext", "keepLines", "pageBreakBefore", "framePr", "widowControl",
            "numPr", "suppressLineNumbers", "pBdr", "shd", "tabs", "suppressAutoHyphens",
            "kinsoku", "wordWrap", "overflowPunct", "topLinePunct", "autoSpaceDE",
            "autoSpaceDN", "bidi", "adjustRightInd", "snapToGrid", "spacing", "ind",
            "contextualSpacing", "mirrorIndents", "suppressOverlap", "jc", "textDirection",
            "textAlignment", "textboxTightWrap", "outlineLvl", "divId", "cnfStyle", "rPr",
            "sectPr", "pPrChange"],
    "tblPr": TBLPR_ORDER,
    "trPr": ["cnfStyle", "divId", "gridBefore", "gridAfter", "wBefore", "wAfter", "cantSplit",
             "trHeight", "tblHeader", "tblCellSpacing", "jc", "hidden", "ins", "del",
             "trPrChange"],
    "tcPr": ["cnfStyle", "tcW", "gridSpan", "hMerge", "vMerge", "tcBorders", "shd", "noWrap",
             "tcMar", "textDirection", "tcFitText", "vAlign", "hideMark", "headers",
             "cellIns", "cellDel", "cellMerge", "tcPrChange"],
    "tcBorders": ["top", "start", "left", "bottom", "end", "right", "insideH", "insideV",
                  "tl2br", "tr2bl"],
}


def validate_docx(path: Path) -> None:
    import zipfile
    from lxml import etree

    w = "{" + W_NS + "}"
    root = etree.fromstring(zipfile.ZipFile(path).read("word/document.xml"))
    problems = []

    def local(el):
        return el.tag.split("}")[1]

    for el in root.iter(w + "p"):
        for child in el:
            if child.tag.startswith("{" + M_NS):
                continue
            if local(child) not in _P_OK:
                problems.append(f"illegal child of w:p: {child.tag}")
        names = [local(c) for c in el]
        if "pPr" in names and names.index("pPr") != 0:
            problems.append(f"w:pPr not first: {names}")
    for el in root.iter(w + "r"):
        for child in el:
            if local(child) not in _R_OK:
                problems.append(f"illegal child of w:r: {child.tag}")
        names = [local(c) for c in el]
        if "rPr" in names and names.index("rPr") != 0:
            problems.append(f"w:rPr not first: {names}")
    for tag, order in _PR_ORDERS.items():
        for el in root.iter(w + tag):
            ranks = [order.index(local(c)) for c in el if local(c) in order]
            if ranks != sorted(ranks):
                problems.append(f"w:{tag} children out of schema order: "
                                f"{[local(c) for c in el]}")
    depth = 0
    for fc in root.iter(w + "fldChar"):
        depth += {"begin": 1, "end": -1}.get(fc.get(w + "fldCharType"), 0)
        if depth < 0:
            problems.append("unbalanced field: an end without a begin")
    if depth != 0:
        problems.append(f"unbalanced field characters (depth {depth})")

    if problems:
        for p in dict.fromkeys(problems):
            print("  INVALID:", p)
        raise SystemExit(f"{len(problems)} OOXML problem(s) – Word would refuse this file.")
    print("Structure check: OK")


# =========================================================================
def main() -> None:
    if DOCX_PATH.exists():
        print(f"Backup of current draft: {backup(DOCX_PATH)}")
    shutil.copy(TEMPLATE_PATH, DOCX_PATH)
    doc = docx.Document(str(DOCX_PATH))

    # The template ships four Heading 1s; the paper needs five sections with
    # subsections under Methodology, so the body between the first heading and
    # the tail anchor is cleared wholesale and rebuilt. The tail anchor carries
    # the 2-column sectPr and must survive.
    h_intro = next(p for p in doc.paragraphs if p.style.name == "Heading 1")
    tail, base_sectpr = find_tail_anchor(doc)
    base_sectpr = copy.deepcopy(base_sectpr)

    # title + author block
    title_p = doc.paragraphs[0]
    for r in list(title_p.runs)[1:]:
        r._r.getparent().remove(r._r)
    title_p.runs[0].text = TITLE
    tighten_author_block(doc, h_intro)
    tighten_styles(doc)
    fix_equation_font(doc)

    clear_between(doc, h_intro, tail)

    omml = json.loads(OMML_PATH.read_text(encoding="utf-8"))

    def full_width(anchor, build):
        section_break(anchor, base_sectpr, 2, 360)      # close the 2-column run
        build(anchor)
        section_break(anchor, base_sectpr, 1, 720)      # island spans the page

    # ---- Abstract and index terms, above the first heading -----------------
    abstract_p = h_intro.insert_paragraph_before(style="Abstract")
    abstract_p.add_run("Abstract—").bold = True
    abstract_p.add_run(ABSTRACT)
    kw_p = h_intro.insert_paragraph_before(style="Keywords")
    kw_p.add_run("Index Terms—").bold = True
    kw_p.add_run(KEYWORDS)

    h_intro.text = "Introduction"
    a = tail          # everything below is inserted before the tail anchor

    # ================= I. INTRODUCTION ===================================
    body(a, [
        "Network businesses are beginning to use retail tariffs to change when "
        "electricity is used. In New South Wales, the Solar Sharer Offer gives "
        "households free electricity in the middle of the day to pull "
        "discretionary consumption into the solar peak. Before rollout its effect "
        "has to be estimated: what the load profile becomes, where the new peak "
        "sits, and how deep the midday minimum goes [[palensky2011]].",

        "That estimate depends on one quantity the tariff itself cannot supply – "
        "how much of the load at each time of day is genuinely able to move. At "
        "network scale the answer must come from data the utility already holds, "
        "and smart meters record only the total energy drawn in each interval: "
        "not which appliances drew it, and not which part could have been "
        "shifted [[kwac2014|wang2019]]. "
        "Fig. {{fig:decomp}} shows the resulting difficulty at a single time of "
        "day. On most days the reading reflects ordinary, non-shiftable activity; "
        "on a minority it also contains a flexible event such as an "
        "electric-vehicle charge or a pool-pump cycle. We call the persistent part "
        "the backbone and the additional part the candidate flexibility to be sized. The "
        "meter reports only their sum, and nothing in the record says which days "
        "contained an event.",
    ])
    picture(a, "decomp", "fig1_decomposition.png",
            "Load at one fixed time of day across forty days. On event days "
            "(markers) a flexible load adds to the backbone $B$; the meter reports "
            "only the total, and a high reading is not by itself evidence of an "
            "event.", COL_W)
    body(a, [
        "In operational models the shiftable share is assumed rather than "
        "measured. In ours – PyNTLP (Python for Network Tariff to Load Profile) – "
        "the energy shifted into the free window on a given day is estimated as",
    ])
    equation(a, omml[8], "mvp")
    body(a, [
        "where E_day is the day's baseline energy (the only measured term), $s$ is a "
        "hand-picked shiftable fraction of that energy, $k$ is an assumed "
        "response-intensity factor, and $a$ is the assumed adoption rate among "
        "eligible customers. This paper addresses $s$. One value is applied to every "
        "customer and every time of day, it says nothing about whether flexible "
        "load is even visible in the data, and the modelled impact scales linearly "
        "with it – so its error passes straight into the network case.",

        "Alternatives exist, but none of the standard ones fits the setting. "
        "Bottom-up appliance models "
        "are physically faithful but need ownership and usage data that is rarely "
        "available at network scale [[wang2019]], and individual switching "
        "patterns smooth out "
        "under aggregation. Non-intrusive load monitoring needs sub-metered "
        "training data [[hart1992]], while customer-baseline methods estimate a "
        "counterfactual for a known event window and so presuppose the event times "
        "[[mathieu2011|valentini2022|zhang2016]].",

        "What is needed is something between the two: practical enough to run from "
        "aggregate meter data alone, yet more physically representative than a "
        "single constant. Fig. {{fig:positioning}} places the options on those two "
        "axes.",
    ])
    picture(a, "positioning", "fig2_positioning.png",
            "Alternative methods on physical representativeness against "
            "practicality; the proposed method keeps the data requirement of a "
            "quantile while removing the arbitrary choice.", COL_W)
    body(a, [
        "The pragmatic compromise is to take a low quantile of the across-day "
        "distribution as the backbone. It needs nothing beyond the aggregate "
        "profile. Its weakness is that the quantile is hand-picked – a convention "
        "rather than an estimate – and the same value is then applied at every "
        "time of day.",

        "This is the problem the paper addresses. The quantile that actually lands "
        "on the backbone is not a constant. It depends on how often flexible "
        "events occur at that timestamp and on how large they are relative to "
        "ordinary day-to-day variation, and both change through the day. A "
        "hand-picked quantile is therefore correct only under a narrow set of "
        "conditions, and carries an uncontrolled – and invisible – error "
        "everywhere else.",

        "The contribution of this paper is to stop hand-picking. Rather than "
        "choosing a quantile to split the load, we derive "
        "the right quantile from the data at each timestamp – a method we call "
        "Adaptive Quantile Flexibility (AQF) – declare when the data "
        "cannot support that estimate, and map where the split is recoverable at "
        "all. The quantile comes from a closed-form relation between backbone "
        "bias, event frequency and event size, the last two estimated per "
        "timestamp by fitting a two-component mixture; a separation diagnostic "
        "blends the estimate towards a default quantile when the fit cannot be "
        "trusted; and the result is capped at the median. The method is evaluated "
        "on synthetic profiles in which the true split is known by construction, "
        "against fixed quantiles at three conventional values.",
    ])

    # ================= II. METHODOLOGY ===================================
    heading(a, "Methodology")
    subheading(a, "Problem formulation")
    body(a, [
        "Fix one time of day – a timestamp, say 18:30 – and consider the load "
        "recorded at that timestamp on each of $D$ days. Writing $L$ for that "
        "reading, we model it as",
    ])
    equation(a, omml[0], "model")
    body(a, [
        "where $B$ is the backbone level, $Z$ is a Bernoulli indicator of whether a "
        "flexible event occurred on that day, $A$ is the event magnitude, and $ε$ is "
        "ordinary "
        "variation with standard deviation $σ$. Events occur independently with "
        "probability $p$, the event frequency at that timestamp. The across-day "
        "distribution of $L$ is therefore a mixture of two Gaussians: a non-event "
        "component centred on $B$ with weight 1−$p$, and an event component centred on "
        "$B$+$A$ with weight $p$.",

        "Two derived quantities govern what follows: the event size $κ$ = $A$/$σ$, the "
        "event magnitude relative to ordinary variation, and the backbone bias "
        "$b$ = ($B̂$ − $B$)/$σ$, the error of a backbone estimate $B̂$ in units of $σ$. It is $κ$, not $p$, that decides whether the two "
        "components can be told apart. Fig. {{fig:mechanism}} makes this concrete: "
        "raising $κ$ pushes the two modes apart until the distribution is visibly "
        "bimodal, whereas raising $p$ merely moves mass into the event mode. If the "
        "backbone is estimated as the $q$-quantile of the observed loads "
        "[[koenker1978]], the bias satisfies",
    ])
    equation(a, omml[2], "relation")
    body(a, [
        "where Φ is the standard normal distribution function. Equation "
        "({{eq:relation}}) is the central relation of the paper: for given $p$ and "
        "$κ$, exactly one quantile places the estimate on $B$, so a fixed quantile is "
        "unbiased only along a one-dimensional curve in the ($p$, $κ$) plane.",
    ])
    full_width(a, lambda anc: picture(
        anc, "mechanism", "fig3_mechanism.png",
        "Across-day load distribution at one timestamp, and the backbone each rule "
        "selects. (a) to (b): event size $κ$ grows at fixed frequency, separating the "
        "modes. (b) to (c): frequency $p$ grows at fixed $κ$, enlarging the event mode "
        "without separating it. The fixed quantile drifts either side of $B$; the "
        "adaptive quantile tracks it.", FULL_W))

    subheading(a, "Adaptive Quantile Flexibility")
    body(a, [
        "Setting $b$ = 0 in ({{eq:relation}}) and solving for $q$ gives the "
        "bias-minimising quantile at timestamp $t$, in which p̂_t and κ̂_t denote "
        "per-timestamp estimates of $p$ and $κ$,",
    ])
    equation(a, omml[3], "qstar")
    body(a, [
        "The correct quantile moves continuously between one half when events are "
        "absent and (1−$p$)/2 when they are fully separated – precisely what a fixed "
        "choice cannot follow. In practice $p$ and $κ$ are unknown, so AQF estimates them: at each timestamp "
        "we fit a two-component Gaussian mixture with tied variance to the "
        "across-day readings by expectation-maximisation [[dempster1977|mclachlan2000]], "
        "using scikit-learn [[pedregosa2011]]. The lower-mean component is taken as "
        "the non-event state, so the upper component's weight estimates $p$ and the "
        "standardised separation of the means estimates $κ$.",

        "A mixture returns parameters whether or not its two components are "
        "genuinely distinguishable, so the estimate is only as trustworthy as that "
        "separation. AQF therefore gates on the standardised separation of the "
        "fitted means, D̂_t = κ̂_t/√2 (not the day count $D$), requiring "
        "$D̂$ ≥ D_th = 2 – a bimodality bar in the spirit of Ashman's D "
        "[[ashman1994]] – which puts the identifiable region at $κ$ ≥ 2.83. "
        "Rather than switching abruptly at that threshold, the estimated quantile "
        "is blended towards a default quantile with weight w_t = min(D̂_t/D_th, 1),",
    ])
    equation(a, omml[5], "guard")
    body(a, [
        "with q_def = 0.2. The blended value is finally clipped at 0.50, so no "
        "timestamp can ever be assigned a quantile above the median however the "
        "mixture behaves. A timestamp whose components are well separated uses its "
        "own estimate; one whose mixture is indistinguishable falls back to the "
        "default; intermediate cases interpolate. Candidate "
        "flexibility on each day is then the non-negative excursion of the load "
        "above the estimated backbone. Derivations and implementation are given in "
        "full in the accompanying repository (github.com/​mssamhan31/​PyNTLP).",
    ])

    subheading(a, "Experimental design")
    body(a, [
        "Ground truth for the flexible share does not exist in the aggregate "
        "meter data utilities hold – that absence is the problem itself – so AQF "
        "is evaluated on synthetic data, in "
        "which the backbone, the event magnitude and the event indicators are all "
        "known by construction. For each scenario we draw $D$ = 365 days from "
        "({{eq:model}}), estimate the backbone with each estimator, score the "
        "residual against the known flexible energy, and aggregate over repeats.",

        "Scenarios span ten event frequencies from 0.05 to 0.95 and ten event sizes "
        "from 0.5 to 5.0, with 20 seeded repeats in each of the 100 combinations – "
        "10,000 estimator runs, each with $B$ = 10 kW and $σ$ = 1 kW. The ranges span "
        "values reported for domestic "
        "appliances in the UK-DALE and REFIT appliance-level datasets and the "
        "National Renewable Energy Laboratory (NREL) end-use load profiles "
        "[[kelly2015|murray2017|wilson2022]]. Five estimators are compared: fixed "
        "quantiles at 0.1, 0.2 and 0.3; oracle-q, given the true $p$ and $κ$; and AQF "
        "itself, which sees only the load. Oracle-q isolates the cost of estimating "
        "$p$ and $κ$, but it is not a performance ceiling: it still takes an empirical "
        "quantile of $D$ days. Two metrics are reported,",
    ])
    equation(a, omml[7], "metrics")
    body(a, [
        "The recovery ratio R_F is estimated flexible energy divided by true "
        "flexible energy, so 1 is perfect recovery, above 1 over-recovery and "
        "below 1 under-recovery; accuracy is summarised as the absolute "
        "deviation |R_F − 1|. MAE_B is the mean absolute error of the backbone "
        "level, averaged over the $N$ scored estimates – one per run, each at "
        "its own timestamp $t$; unlike R_F it is unaffected by the non-negative "
        "truncation introduced in Section II-B.",
    ])

    # ================= III. RESULTS ======================================
    heading(a, "Results")
    build_results_summary_table()   # keeps the legacy CSV current; not printed
    t3, t3_headers, t3_bold, t3_italic = build_headline_table()
    table_caption(a, "results",
                  "Estimator accuracy over the 100 combinations of event frequency "
                  "and event size. Best deployable value per column in bold. "
                  "Oracle-q (italic) is given the true event frequency and event "
                  "size; it is a reference, not an available baseline and not a "
                  "performance ceiling.")
    dataframe_table(
        doc, a, t3,
        [Inches(0.86), Inches(0.60), Inches(0.62), Inches(0.60), Inches(0.62)],
        align=["l", "c", "c", "c", "c"], headers=t3_headers,
        bold_cells=t3_bold, italic_rows=t3_italic)
    body(a, [
        "Table {{tab:results}} reports accuracy over the grid. Ranked by the mean "
        "of |R_F − 1|, AQF at 1.54 and the best fixed quantile at 1.55 look "
        "indistinguishable, but that ranking is an artefact: the mean is dominated "
        "by the rare-event cells where R_F exceeds 15 for every estimator, "
        "including oracle-q, for a reason established below that is not the "
        "estimator's doing. On the median, which that tail does not dominate, AQF "
        "deviates by 0.26 against 0.60 for the best fixed quantile, a factor of "
        "2.3. On backbone error, which the effect does not touch at all, AQF "
        "reaches 0.26 kW against 0.75 kW for the best fixed quantile on that "
        "metric ($q$ = 0.2).",

        "The identifiable region sharpens the comparison: there the median deviation "
        "is 0.114 for AQF against 0.105 for oracle-q, while the best fixed "
        "quantile on that metric ($q$ = 0.2) manages only 0.538. Restricting "
        "further to event frequencies of 0.3 and above gives 0.0659 against "
        "0.0657: within the identifiable region, at moderate-to-high event "
        "frequency, AQF matches a method given the true parameters to within 0.0002. "
        "The identifiability test is therefore not a disclaimer but a statement "
        "of when the output can be relied upon.",

        "Fig. {{fig:map}} shows where each estimator works. The fixed quantile is "
        "unbiased only along a narrow contour and under-recovers badly when events "
        "are frequent and large, where a quantile fixed at 0.3 cuts into the event "
        "mode itself; AQF removes almost all of that region. Cell by cell, AQF "
        "beats fixed quantiles of 0.1, 0.2 and 0.3 in 95, 89 and 58 of the 100 "
        "cells, and a per-cell hindsight choice in 44 – reported deliberately, "
        "since deploying that choice requires knowing what AQF estimates.",
    ])
    full_width(a, lambda anc: picture(
        anc, "map", "fig4_recoverability_map.png",
        "Recovery ratio R_F over the grid for (a) the best fixed quantile, (b) "
        "AQF and (c) oracle-q. White is perfect recovery, red over- and blue "
        "under-recovery on a log₂ scale capped at eight; the contour marks R_F = 1 "
        "and the dashed line the identifiability threshold.", FULL_W))
    picture(a, "curves", "fig5_curves.png",
            "(a) Median |R_F − 1| and (b) mean backbone error against event size, "
            "shaded where the diagnostic clears its threshold. (c) Recovery ratio "
            "against event frequency at $κ$ = 3, with the parameter-free noise-floor "
            "prediction (dotted).", COL_W)
    body(a, [
        "Fig. {{fig:curves}} explains why. Fixed-quantile accuracy plateaus in $κ$: a "
        "constant cannot exploit a clearer signal, whereas AQF converges onto "
        "oracle-q – and above $κ$ ≈ 4 overtakes it on backbone error, consistent "
        "with a fitted curve avoiding sampling noise an empirical quantile "
        "carries even given the true parameters.",

        "The remaining error at low event frequency is not estimator bias: every "
        "non-event day adds $σ$/√(2π) ≈ 0.399$σ$ of spurious flexibility even with an "
        "exact backbone, predicting the noise floor R_F = 1 + (1−$p$)/($pκ$√(2π)) "
        "with no fitted "
        "parameter. Oracle-q follows it to a median ratio of 1.006, so those large "
        "ratios are a property of the truncated-residual estimator.",
    ])

    # ================= IV. DISCUSSION AND CONCLUSION =====================
    heading(a, "Discussion and Conclusion")
    body(a, [
        "The quantile a fixed-quantile estimator applies is a convention, not an "
        "estimate, and its bias varies with quantities the analyst never sees. This "
        "paper derived the relation those quantities satisfy and replaced the "
        "convention with a per-timestamp estimate, gated by a diagnostic that "
        "reports when it should not be trusted and capped at the median. Three "
        "limitations bound the claim: good recovery shows identifiability of a "
        "shape, not physical shiftability; the evaluation is synthetic; and the "
        "noise floor bounds any truncated-residual estimator. Next steps are to "
        "validate the estimated $p$ and $κ$ against UK-DALE and REFIT ground truth "
        "[[kelly2015|murray2017]], then replace the hand-picked fraction $s$ in "
        "({{eq:mvp}}) with the calibrated per-timestamp estimate.",
    ])

    # ================= REFERENCES ========================================
    h_ref = a.insert_paragraph_before("References", style="Heading 1")
    unnumber(h_ref)
    # IEEE order: as first cited in the text. Anything defined but never cited
    # would silently vanish from the list, so check that explicitly.
    uncited = [k for k in refs.REFS if k not in _cite_order]
    if uncited:
        raise SystemExit(f"references defined but never cited: {sorted(uncited)}")
    bib_paras = []
    for n, key in enumerate(_cite_order, start=1):
        p = a.insert_paragraph_before(style="references")
        p.add_run(f"[{n}]\t")
        for seg, ital in refs.ieee_entry_parts(key):
            p.add_run(seg).italic = ital
        unnumber(p)
        p.paragraph_format.left_indent = Inches(0.22)
        p.paragraph_format.first_line_indent = Inches(-0.22)   # number hangs left
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        bib_paras.append(p)
    # wrap the whole list in one Zotero bibliography field
    bib_instr = (' ADDIN ZOTERO_BIBL {"uncited":[],"omitted":[],"custom":[]} '
                 'CSL_BIBLIOGRAPHY ')
    first, last = bib_paras[0], bib_paras[-1]
    lead = first.runs[0]._r
    for el in (fld_char("begin"), instr_text(bib_instr), fld_char("separate")):
        run = OxmlElement("w:r")      # fldChar/instrText are only legal inside w:r
        run.append(el)
        lead.addprevious(run)         # each lands directly before the text run
    last.add_run()._r.append(fld_char("end"))

    # Zotero document preferences, parked in the trailing paragraph. The style is
    # IEEE so that refreshing the fields in Zotero reproduces what is written here.
    prefs = doc.paragraphs[-1]
    # This paragraph is structural – it carries the final 2-column sectPr and
    # cannot be deleted – but at full body size its empty mark alone spills onto
    # a fifth page. Collapse it to 1 pt with no leading.
    set_mark_size(prefs, 2)
    kill_autospacing(prefs)
    prefs_instr = (
        ' ADDIN ZOTERO_PREF_1 <data data-version="3" zotero-version="6.0.36">'
        '<session id="aqfIsgt26"/>'
        '<style id="http://www.zotero.org/styles/vancouver" locale="en-US" hasBibliography="1" '
        'bibliographyStyleHasBeenSet="1"/><prefs>'
        '<pref name="fieldType" value="Field"/>'
        '<pref name="automaticJournalAbbreviations" value="false"/>'
        '<pref name="noteType" value="0"/></prefs></data> '
    )
    prefs.add_run()._r.append(fld_char("begin"))
    prefs.add_run()._r.append(instr_text(prefs_instr))
    prefs.add_run()._r.append(fld_char("separate"))
    prefs.add_run()._r.append(fld_char("end"))

    doc.save(str(DOCX_PATH))
    validate_docx(DOCX_PATH)
    refs.write_ris(RIS_PATH)
    print(f"Rebuilt: {DOCX_PATH}")
    print(f"RIS exported: {RIS_PATH}")


if __name__ == "__main__":
    main()
